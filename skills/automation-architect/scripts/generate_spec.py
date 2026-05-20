#!/usr/bin/env python3
"""
Generate a formatted Strategic AI Operating System Spec document.

Input JSON schema (spec_data.json):
{
    "system_name": "Email Reply Agent",
    "generated_date": "Monday, May 4, 2026",

    "phase1": {
        "target_user": "Solo founder handling 50+ emails/day",
        "core_outcome": "Inbox zero without reading every thread",
        "problem_solved": "Spending 2hrs/day triaging and writing email replies",
        "success_criteria": "Drafts ready in inbox within 5 minutes, saves 1.5hrs/day"
    },

    "phase2": {
        "trigger": "New email arrives in Gmail inbox",
        "input": "Email thread content (subject, body, sender, history)",
        "components": [
            {"name": "Email Parser", "role": "Extracts intent, urgency, and required action"},
            {"name": "Context Retriever", "role": "Pulls relevant prior email history"},
            {"name": "Reply Drafter", "role": "Generates context-aware draft reply"}
        ],
        "output": "Draft reply saved to Gmail drafts folder",
        "agent_roles": "N/A — linear workflow",
        "data_flow": "Email thread → parsed intent → draft reply → Gmail drafts"
    },

    "phase3": {
        "selected_tool": "Make",
        "justification": "Multi-step workflow with Gmail integration, data transformation, and Claude API call — Make handles this cleanly at low cost.",
        "rejected_tools": [
            {"tool": "Zapier", "reason": "Too expensive at daily volume, limited data transformation"},
            {"tool": "Gumloop", "reason": "Less mature Gmail integration for read/write flows"}
        ],
        "documentation_url": "https://developers.make.com/",
        "setup_start": "make.com → New scenario → Gmail trigger module → HTTP (Claude API) → Gmail create draft"
    },

    "phase4": {
        "steps": [
            {"step": 1, "input": "New Gmail email", "action": "Trigger: watch for new emails", "output": "Email object"},
            {"step": 2, "input": "Email object", "action": "Parse subject, body, sender", "output": "Structured intent data"},
            {"step": 3, "input": "Intent data", "action": "Send to Claude API with system prompt", "output": "Draft reply text"},
            {"step": 4, "input": "Draft reply text", "action": "Create Gmail draft", "output": "Draft saved in inbox"}
        ],
        "error_handling": "If Claude API fails, log error and skip — do not send partial drafts",
        "estimated_time": "15–30 seconds per email"
    },

    "phase5_steps": [
        "Create Make account at make.com",
        "Add Gmail module: Watch emails (trigger)",
        "Add Text Parser module: extract subject and body",
        "Add HTTP module: POST to Claude API (claude-sonnet-4-6)",
        "Configure system prompt for email context",
        "Add Gmail module: Create draft with API response",
        "Test with one real email",
        "Activate scenario"
    ],

    "phase6": {
        "bottleneck": "Claude API response time (~5–8s) is the slowest step",
        "quick_fix": "Use claude-haiku-4-5 for speed if reply quality is acceptable",
        "cost_opportunity": "Switch from Make to n8n self-hosted to eliminate monthly SaaS cost",
        "reliability_risk": "Gmail API rate limits at >100 emails/day — add delay buffer",
        "ai_expansion": "Add an intent classifier to skip newsletters and auto-archive them",
        "multi_agent_potential": "Add a second agent to handle follow-up scheduling after reply is sent",
        "next_scaling_step": "Add a Notion knowledge base so the agent pulls from your real docs and past replies"
    }
}
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...", flush=True)
    import subprocess
    subprocess.run(
        [sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages"],
        capture_output=True
    )
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Colors
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
ORANGE     = RGBColor(0xC5, 0x5A, 0x11)
GREEN      = RGBColor(0x1E, 0x7E, 0x34)
MID_GRAY   = RGBColor(0x55, 0x55, 0x55)
BLACK      = RGBColor(0x00, 0x00, 0x00)
FOOTER_CLR = RGBColor(0xAA, 0xAA, 0xAA)
LIGHT_BG   = RGBColor(0xF7, 0xF9, 0xFC)


def _color(run, rgb):
    run.font.color.rgb = rgb


def _font(run, size=11, bold=False, italic=False, name="Calibri"):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic


def _add_rule(doc, color="CCCCCC", thickness="6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper() if level == 1 else text)
    size = 9 if level == 1 else 11
    _font(run, size=size, bold=True)
    _color(run, DARK_BLUE)
    if level == 1:
        _add_rule(doc, color="1F4E79", thickness="8")
    return p


def _body(doc, text, color=None, bold=False, italic=False, size=11, indent=0.0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    _font(run, size=size, bold=bold, italic=italic)
    if color:
        _color(run, color)
    return p


def _labeled(doc, label, value, label_color=None, value_color=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    lr = p.add_run(f"{label}: ")
    _font(lr, size=size, bold=True)
    _color(lr, label_color or DARK_BLUE)
    vr = p.add_run(value)
    _font(vr, size=size)
    _color(vr, value_color or BLACK)
    return p


def _bullet(doc, text, color=None, indent=0.2, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    br = p.add_run("▸  ")
    _font(br, size=size)
    _color(br, ORANGE)
    tr = p.add_run(text)
    _font(tr, size=size)
    _color(tr, color or BLACK)
    return p


def _checkbox(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    cr = p.add_run("☐  ")
    _font(cr, size=size)
    tr = p.add_run(text)
    _font(tr, size=size)
    return p


def generate_document(data: dict, output_path: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    system_name = data.get("system_name", "AI System")
    gen_date = data.get("generated_date", "")

    # ── COVER ──────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run("STRATEGIC AI OPERATING SYSTEM SPEC")
    _font(tr, size=11, bold=True)
    _color(tr, MID_GRAY)

    sys_p = doc.add_paragraph()
    sys_p.paragraph_format.space_after = Pt(2)
    sr = sys_p.add_run(system_name)
    _font(sr, size=26, bold=True)
    _color(sr, DARK_BLUE)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    sbr = sub_p.add_run(f"Generated by Claude Automation Architect  ·  {gen_date}")
    _font(sbr, size=10, italic=True)
    _color(sbr, MID_GRAY)

    _add_rule(doc, color="1F4E79", thickness="12")

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1: STRATEGIC AI OPERATING SYSTEM SPEC
    # ══════════════════════════════════════════════════════════════════════
    _section_header(doc, "Section 1 — Strategic AI Operating System Spec")

    # System Summary
    _section_header(doc, "System Summary", level=2)
    p1 = data.get("phase1", {})
    _labeled(doc, "Target user", p1.get("target_user", ""))
    _labeled(doc, "Core outcome", p1.get("core_outcome", ""))
    _labeled(doc, "Problem solved", p1.get("problem_solved", ""))
    _labeled(doc, "Success criteria", p1.get("success_criteria", ""))

    # Architecture Overview
    _section_header(doc, "Architecture Overview", level=2)
    p2 = data.get("phase2", {})
    _labeled(doc, "Trigger", p2.get("trigger", ""))
    _labeled(doc, "Input", p2.get("input", ""))

    components = p2.get("components", [])
    if components:
        _body(doc, "Components:", bold=True, size=10.5)
        for i, comp in enumerate(components, 1):
            _bullet(doc, f"{comp.get('name', '')} — {comp.get('role', '')}")

    _labeled(doc, "Output", p2.get("output", ""))
    _labeled(doc, "Agent roles", p2.get("agent_roles", "N/A"))
    _labeled(doc, "Data flow", p2.get("data_flow", ""))

    # Tool Rationale
    _section_header(doc, "Tool Rationale", level=2)
    p3 = data.get("phase3", {})
    _labeled(doc, "Selected tool", p3.get("selected_tool", ""), label_color=GREEN)
    _body(doc, p3.get("justification", ""), size=10.5, indent=0.1)

    rejected = p3.get("rejected_tools", [])
    if rejected:
        _body(doc, "Rejected alternatives:", bold=True, size=10.5)
        for item in rejected:
            _bullet(doc, f"{item.get('tool', '')} — {item.get('reason', '')}")

    if p3.get("documentation_url"):
        _labeled(doc, "Documentation", p3["documentation_url"])
    if p3.get("setup_start"):
        _labeled(doc, "Setup start", p3["setup_start"])

    # Workflow Logic
    _section_header(doc, "Workflow Logic", level=2)
    p4 = data.get("phase4", {})
    steps = p4.get("steps", [])
    for step in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.1)
        nr = p.add_run(f"Step {step.get('step', '')}: ")
        _font(nr, size=10.5, bold=True)
        _color(nr, ORANGE)
        ar = p.add_run(f"{step.get('input', '')} → {step.get('action', '')} → {step.get('output', '')}")
        _font(ar, size=10.5)
        _color(ar, BLACK)

    if p4.get("error_handling"):
        _labeled(doc, "Error handling", p4["error_handling"])
    if p4.get("estimated_time"):
        _labeled(doc, "Est. execution time", p4["estimated_time"])

    # Scaling Strategy
    _section_header(doc, "Scaling Strategy", level=2)
    p6 = data.get("phase6", {})
    _labeled(doc, "Bottleneck", p6.get("bottleneck", ""))
    _labeled(doc, "Quick fix", p6.get("quick_fix", ""))
    _labeled(doc, "Reliability risk", p6.get("reliability_risk", ""))
    _labeled(doc, "AI expansion", p6.get("ai_expansion", ""))
    _labeled(doc, "Multi-agent potential", p6.get("multi_agent_potential", ""))
    _labeled(doc, "Next scaling step", p6.get("next_scaling_step", ""), label_color=GREEN)

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2: IMPLEMENTATION BLUEPRINT
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _section_header(doc, "Section 2 — Implementation Blueprint")

    # Build Instructions
    _section_header(doc, "Step-by-Step Build Instructions", level=2)
    p5_steps = data.get("phase5_steps", [])
    for i, step_text in enumerate(p5_steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.1)
        nr = p.add_run(f"{i}.  ")
        _font(nr, size=10.5, bold=True)
        _color(nr, DARK_BLUE)
        tr2 = p.add_run(step_text)
        _font(tr2, size=10.5)

    # Tool Setup Guide
    _section_header(doc, "Tool Setup Guide", level=2)
    selected_tool = p3.get("selected_tool", "")
    _body(doc, f"Primary tool: {selected_tool}", bold=True, size=10.5)
    _body(doc, f"Documentation: {p3.get('documentation_url', '')}", size=10.5, indent=0.1)
    _body(doc, f"Where to start: {p3.get('setup_start', '')}", size=10.5, indent=0.1)

    # Execution Checklist
    _section_header(doc, "Execution Checklist", level=2)
    _body(doc, "Use this to track your build progress:", italic=True, size=10, color=MID_GRAY)
    for step in steps:
        _checkbox(doc, f"Step {step.get('step', '')}: {step.get('action', '')}")
    for p5_step in p5_steps:
        _checkbox(doc, p5_step)

    # Footer
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_p.paragraph_format.space_before = Pt(20)
    fr = foot_p.add_run(f"Generated by Claude Automation Architect  ·  {gen_date}  ·  Full Agent Builder OS")
    _font(fr, size=9, italic=True)
    _color(fr, FOOTER_CLR)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"✓ Saved: {out}")
    return str(out)


def main():
    parser = argparse.ArgumentParser(description="Generate AI OS Spec .docx")
    parser.add_argument("--input",  required=True, help="Path to spec_data.json")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    with open(args.input, "r") as f:
        data = json.load(f)

    result = generate_document(data, args.output)
    print(f"Document ready: {result}")


if __name__ == "__main__":
    main()
